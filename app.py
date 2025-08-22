import streamlit as st
from docx import Document
import io
from datetime import datetime
import zipfile

# Replaces placeholders in a paragraph while preserving the style of the first run
def replace_placeholders_in_paragraph(paragraph, values):
    full_text = ''.join(run.text for run in paragraph.runs)
    for key, val in values.items():
        placeholder = f"{{{{{key}}}}}"
        full_text = full_text.replace(placeholder, str(val))

    if full_text != paragraph.text:
        if paragraph.runs:
            first_run = paragraph.runs[0]
            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic
            new_run.underline = first_run.underline
            new_run.font.name = first_run.font.name
            new_run.font.size = first_run.font.size

def replace_placeholders_in_doc(doc, values):
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, values)
    return doc

# --- Streamlit UI ---
st.set_page_config(page_title="Contract Generator", layout="centered")
st.title("📄 Multi-Contract Generator")

st.markdown("Upload multiple Word `.docx` templates with placeholders like `{{Date}}`, `{{contract_name}}`, etc.")

uploaded_files = st.file_uploader("Upload templates (.docx)", type="docx", accept_multiple_files=True)

with st.form("input_form"):
    st.subheader("📝 Fill in shared placeholder values")

    contract_name = st.text_input("Contract Name", "")
    invoice_no = st.text_input("Invoice No", "")
    amount = st.text_input("Amount", "")
    weight = st.text_input("Weight", "")
    bl = st.text_input("BL Number", "")
    date = st.text_input("Date", datetime.today().strftime("%B %d, %Y"))
    Buyer = st.text_input("Buyer", "")
    Address = st.text_input("Address", "")
    IEC = st.text_input("IEC CODE", "")
    Email = st.text_input("Email-id", "")

    submitted = st.form_submit_button("Generate Contracts")

if submitted and uploaded_files:
    shared_values = {
    "Date": "",
    "invoice_no": "",
    "contract_name": "",
    "amount": "",
    "weight": "",
    "BL": "",
    "Buyer": "",
    "Address": "",
    "IEC": "",
    "Email-id": ""
    }

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zip_file:
        for uploaded_file in uploaded_files:
            doc = Document(uploaded_file)
            modified_doc = replace_placeholders_in_doc(doc, shared_values)

            output_buffer = io.BytesIO()
            output_name = f"{uploaded_file.name.replace('.docx', '')}_generated.docx"
            modified_doc.save(output_buffer)
            zip_file.writestr(output_name, output_buffer.getvalue())

    st.success("✅ Contracts generated!")

    st.download_button(
        label="📥 Download All as ZIP",
        data=zip_buffer.getvalue(),
        file_name="generated_contracts.zip",
        mime="application/zip"
    )

elif submitted:

    st.warning("⚠️ Please upload at least one template file.")






