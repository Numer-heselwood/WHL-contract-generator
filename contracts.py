from docx import Document

def replace_placeholders_in_paragraph(paragraph, values):
    # Step 1: Get full paragraph text
    full_text = ''.join(run.text for run in paragraph.runs)
    
    # Step 2: Replace placeholders
    for key, val in values.items():
        placeholder = f"{{{{{key}}}}}"
        full_text = full_text.replace(placeholder, str(val))
    
    # Step 3: If anything changed, update paragraph
    if full_text != paragraph.text:
        # Use formatting of first run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic
            new_run.underline = first_run.underline
            new_run.font.name = first_run.font.name
            new_run.font.size = first_run.font.size

def generate_contract(template_path, output_path, values):
    doc = Document(template_path)

    # Replace in body
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, values)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, values)

    doc.save(output_path)
    print(f"✅ Contract generated and saved as: {output_path}")
# Your shared values
shared_values = {
    "Date": "August 4, 2025",
    "invoice_no": "C1-2025-000006",
    "contract_name": "SC-2025-000005",
    "amount": "£50000",
    "weight": "500.60MT",
    "BL": "232456446"
}

# Your templates and outputs
contracts_to_generate = [
    {
        "template": r"C:\Users\numer\OneDrive\Desktop\contract generator\origin_template.docx",
        "output":   r"C:\Users\numer\OneDrive\Desktop\contract generator\certificate_of_origin.docx"
    },
    {
        "template": r"C:\Users\numer\OneDrive\Desktop\contract generator\weight_template.docx",
        "output":   r"C:\Users\numer\OneDrive\Desktop\contract generator\certificate_of_weight.docx"
    },
    {
        "template": r"C:\Users\numer\OneDrive\Desktop\contract generator\nowar_template.docx",
        "output":   r"C:\Users\numer\OneDrive\Desktop\contract generator\certificate_of_no_war.docx"
    },
    {
        "template": r"C:\Users\numer\OneDrive\Desktop\contract generator\analysis_template.docx",
        "output":   r"C:\Users\numer\OneDrive\Desktop\contract generator\certificate_of_analysis.docx"
    }
]

# Generate all contracts
for contract in contracts_to_generate:
    generate_contract(contract["template"], contract["output"], shared_values)